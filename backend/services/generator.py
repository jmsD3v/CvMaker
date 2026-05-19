import os
import json
import io
import google.generativeai as genai
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from models.schemas import UserProfile

genai.configure(api_key=os.getenv("GEMINI_API_KEY", ""))
model = genai.GenerativeModel("gemini-2.0-flash-exp")

CV_PROMPT = """You are an expert ATS-optimized CV writer in Spanish. Generate a tailored CV for this candidate.

CANDIDATE PROFILE:
{profile_json}

JOB POSTING:
{job_posting}

INSTRUCTIONS:
- Rank experience and certifications by relevance to this specific job
- Extract and inject exact keywords from the job posting naturally
- ATS format: no tables, no columns, plain linear text
- Summary: 3-4 sentences in Spanish highlighting strongest match points
- Experience bullets: action verb + metric/result format, in Spanish
- Certifications: list only those relevant, most relevant first
- Skills: extract from profile, prioritize what the job mentions

Return ONLY valid JSON (no markdown, no explanation):
{{
  "summary": "...",
  "experience": [
    {{"title": "...", "company": "...", "period": "...", "bullets": ["...", "..."]}}
  ],
  "education": [
    {{"degree": "...", "institution": "...", "period": "..."}}
  ],
  "certifications": [
    {{"name": "...", "issuer": "...", "date": "..."}}
  ],
  "skills": ["...", "..."]
}}"""

LETTER_PROMPT = """Escribí una carta de presentación profesional en español para este candidato.

CANDIDATO: {full_name}
CONTACTO: {contact}
RESUMEN CV: {summary}
AVISO LABORAL: {job_posting}

La carta debe:
- Tener 3-4 párrafos
- Mencionar detalles específicos del aviso laboral
- Destacar las coincidencias más fuertes entre candidato y puesto
- Terminar con un llamado a la acción claro
- Tono: profesional, directo, seguro

Devolvé solo el texto de la carta, sin asunto, sin JSON:"""


def build_cv_prompt(profile: UserProfile, job_posting: str) -> str:
    return CV_PROMPT.format(
        profile_json=profile.model_dump_json(indent=2),
        job_posting=job_posting[:3000],
    )


def generate_cv_content(profile: UserProfile, job_posting: str) -> dict:
    prompt = build_cv_prompt(profile, job_posting)
    response = model.generate_content(prompt)
    raw = response.text.strip()
    if raw.startswith("```"):
        parts = raw.split("```")
        raw = parts[1] if len(parts) > 1 else raw
        if raw.startswith("json"):
            raw = raw[4:]
    return json.loads(raw.strip())


def generate_cover_letter(profile: UserProfile, job_posting: str, cv_summary: str) -> str:
    contact_parts = [x for x in [profile.email, profile.phone, profile.linkedin] if x]
    contact = " | ".join(contact_parts)
    prompt = LETTER_PROMPT.format(
        full_name=profile.full_name,
        contact=contact,
        summary=cv_summary,
        job_posting=job_posting[:2000],
    )
    response = model.generate_content(prompt)
    return response.text.strip()


def _add_section_heading(doc: Document, text: str):
    p = doc.add_heading(text, level=1)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
        run.font.size = Pt(12)


def build_docx(cv_content: dict, full_name: str, contact: str) -> bytes:
    doc = Document()

    normal = doc.styles['Normal']
    normal.font.name = 'Calibri'
    normal.font.size = Pt(11)

    title = doc.add_heading(full_name, 0)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    doc.add_paragraph(contact)
    doc.add_paragraph("")

    _add_section_heading(doc, "RESUMEN PROFESIONAL")
    doc.add_paragraph(cv_content.get("summary", ""))

    _add_section_heading(doc, "EXPERIENCIA LABORAL")
    for exp in cv_content.get("experience", []):
        p = doc.add_paragraph()
        run = p.add_run(f"{exp.get('title', '')} | {exp.get('company', '')}")
        run.bold = True
        doc.add_paragraph(exp.get("period", ""))
        for bullet in exp.get("bullets", []):
            doc.add_paragraph(f"• {bullet}")
        doc.add_paragraph("")

    _add_section_heading(doc, "EDUCACIÓN")
    for edu in cv_content.get("education", []):
        p = doc.add_paragraph()
        p.add_run(edu.get("degree", "")).bold = True
        doc.add_paragraph(f"{edu.get('institution', '')} | {edu.get('period', '')}")

    _add_section_heading(doc, "CERTIFICACIONES")
    for cert in cv_content.get("certifications", []):
        doc.add_paragraph(f"• {cert.get('name', '')} — {cert.get('issuer', '')} ({cert.get('date', '')})")

    _add_section_heading(doc, "COMPETENCIAS TÉCNICAS")
    skills = cv_content.get("skills", [])
    doc.add_paragraph(" | ".join(skills))

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def build_letter_docx(letter_text: str, full_name: str) -> bytes:
    doc = Document()
    doc.add_heading(f"Carta de Presentación — {full_name}", 0)
    doc.add_paragraph("")
    for para in letter_text.split("\n\n"):
        if para.strip():
            doc.add_paragraph(para.strip())
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
